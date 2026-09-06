#!/usr/bin/env python3
"""
SINOV — HUJJAT MATNI (P0-2)
===========================
`etl_doc_text.py` va `tender_document_text` jadvalini tekshiradi.

Ishga tushirish:
    python _tests/doctext_test.py            # to'liq (tarmoq bilan)
    python _tests/doctext_test.py --offline  # faqat parserlar (tarmoqsiz)

Uvicorn ISHGA TUSHIRILMAYDI — to'g'ridan-to'g'ri funksiyalar chaqiriladi.
Model/AI chaqiruvi YO'Q.
"""
import argparse
import io
import os
import sys

sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# KONSOL KODLASHI — Windows kod sahifasidan MUSTAQIL UTF-8.
#
# Chiqish QUVUR yoki FAYLGA yo'naltirilganda (ya'ni CI da) Python
# `locale.getpreferredencoding()` ni oladi — bu mashinada `cp1251`.
# O'zbek kirill (`ҳ`, `қ`, `ў`) va to'liq kenglikdagi belgilar
# (`）`) u yerda YO'Q va chop etish `UnicodeEncodeError` bilan
# BUTUN TO'PLAMNI o'ldiradi. `import_test` aynan shu sababdan
# 143 ta tekshiruvni bajarmasdan yiqilardi. Tafsilot: _tests/konsol.py
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
import konsol  # noqa: E402
import rejim  # noqa: E402

konsol.sozla()

from dotenv import load_dotenv  # noqa: E402

load_dotenv()

import io                                           # noqa: E402
import zipfile                                      # noqa: E402

import psycopg2                                     # noqa: E402
from psycopg2.extras import RealDictCursor          # noqa: E402

import etl_doc_text as E                            # noqa: E402

PASS = FAIL = 0


def check(name: str, cond: bool, detail: str = "") -> None:
    global PASS, FAIL
    if cond:
        PASS += 1
        print(f"  OK   {name}" + (f" — {detail}" if detail else ""))
    else:
        FAIL += 1
        print(f"  XATO {name}" + (f" — {detail}" if detail else ""))


def section(title: str) -> None:
    print(f"\n=== {title} ===")


# ---------------------------------------------------------------------------
# 1. SXEMA
# ---------------------------------------------------------------------------
def test_schema(conn) -> None:
    section("1. Sxema")
    with conn.cursor() as cur:
        cur.execute("""SELECT column_name FROM information_schema.columns
                       WHERE table_name = 'tender_document_text'""")
        cols = {r[0] for r in cur.fetchall()}
    need = {"tender_id", "file_ref", "text", "status", "char_count",
            "page_count", "error", "extractor", "extracted_at"}
    check("jadval va ustunlar mavjud", need <= cols,
          f"yetishmayotgan: {sorted(need - cols) or 'yo‘q'}")

    with conn.cursor() as cur:
        cur.execute("""SELECT indexname FROM pg_indexes
                       WHERE tablename = 'tender_document_text'""")
        idx = {r[0] for r in cur.fetchall()}
    check("indekslar o'rnatilgan", len(idx) >= 4, ", ".join(sorted(idx)))


# ---------------------------------------------------------------------------
# 2. PARSERLAR — tarmoqsiz, sun'iy fayllarda (determinizm tekshiruvi)
# ---------------------------------------------------------------------------
MARKER = "Texnik topshiriq: kompressor 10 atm, kafolat 24 oy"


def _make_docx() -> bytes:
    import docx
    d = docx.Document()
    d.add_paragraph(MARKER)
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "Muddat"
    t.rows[0].cells[1].text = "30 kun"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_xlsx() -> bytes:
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Spetsifikatsiya"
    ws.append(["Nomi", "Soni"])
    ws.append([MARKER, 5])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_parsers() -> None:
    section("2. Parserlar (tarmoqsiz)")

    # DOCX — paragraf VA jadval katakchalari
    text, pages, ex, err = E.extract_docx(_make_docx())
    check("DOCX matn ajratildi", MARKER in text and err is None, f"{len(text)} belgi, {ex}")
    check("DOCX jadval katakchalari ham olindi", "30 kun" in text)

    # XLSX — varaq nomi va katakchalar
    text, pages, ex, err = E.extract_xlsx(_make_xlsx())
    check("XLSX matn ajratildi", MARKER in text and err is None, f"{len(text)} belgi, {ex}")
    check("XLSX varaq nomi qo'shildi", "[Spetsifikatsiya]" in text)
    check("XLSX varaq soni", pages == 1, f"page_count={pages}")

    # HTML — teglar tashlanadi, script/style olinmaydi
    html = b"<html><head><style>p{color:red}</style></head><body><p>Talab: ISO 9001</p><script>var x=1</script></body></html>"
    text, _, ex, err = E.extract_plain(html, "html")
    check("HTML matni tozalandi", "ISO 9001" in text and "color:red" not in text
          and "var x" not in text, repr(text[:60]))

    # CSV / TXT + kodlash (cp1251 — manbada ko'p uchraydi)
    text, _, _, _ = E.extract_plain("Талабнома: 5 дона".encode("cp1251"), "txt")
    check("cp1251 kodlash o'qildi", "Талабнома" in text, repr(text[:40]))

    # NUL bayt — PostgreSQL TEXT uni qabul qilmaydi, tozalanishi SHART
    text, _, _, _ = E.extract_plain(b"Talab:\x00 5 dona\x01 kafolat", "txt")
    check("NUL/boshqaruv belgilari tozalandi",
          "\x00" not in text and "\x01" not in text and "5 dona" in text, repr(text))

    # BUZILGAN fayl -> parser xato beradi (unreadable ga aylanadi)
    _, _, _, err = E.extract_docx(b"not a docx at all, just bytes")
    check("buzilgan DOCX xato qaytardi", err is not None, (err or "")[:50])
    _, _, _, err = E.extract_pdf(b"%PDF-1.4 broken")
    check("buzilgan PDF xato qaytardi", err is not None, (err or "")[:50])


# ---------------------------------------------------------------------------
# 3. FORMAT ANIQLASH va STATUS mantig'i (tarmoqsiz)
# ---------------------------------------------------------------------------
def _zip_yasa(azolar) -> bytes:
    """Sinov uchun ZIP quradi — diskka tegmasdan, xotirada."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        for nom, matn in azolar:
            z.writestr(nom, matn)
    return buf.getvalue()


def test_zip() -> None:
    """ZIP arxivlari — HUJUM HOLATLARI va sehrli baytlar.

    Arxiv TASHQI MANBADAN keladi. Tarmoq talab qilinmaydi: barcha
    holatlar shu yerda quriladi.
    """
    section("3b. ZIP arxivlari — xavfsizlik")

    # --- ZIP SLIP -----------------------------------------------------
    # Diskka HECH QACHON yozmaymiz, shuning uchun slip TUZILISHIGA KO'RA
    # mumkin emas. Shunday bo'lsa ham nom yorliq sifatida tozalanadi —
    # javobda chalg'ituvchi yo'l ko'rinmasin.
    matn, _, _, _ = E.extract_zip(
        _zip_yasa([("../../../../etc/passwd.txt", "root:x:0:0")]))
    check("zip slip: yo'l yorliqdan olib tashlanadi",
          "../" not in matn and "passwd.txt" in matn, matn[:80])

    # --- ZIP BOMBA ----------------------------------------------------
    # Chegara OCHISHDAN OLDIN, e'lon qilingan siqilmagan hajm bo'yicha.
    matn, _, _, xato = E.extract_zip(
        _zip_yasa([("katta.txt", "A" * (E.ZIP_MAX_MEMBER + 1000))]))
    check("zip bomba: katta a'zo o'tkazilmaydi",
          bool(xato) or "A" * 100 not in matn, f"xato={xato}")

    matn, _, _, xato = E.extract_zip(
        _zip_yasa([(f"f{i}.txt", "matn " * 20)
                   for i in range(E.ZIP_MAX_MEMBERS + 5)]))
    check("a'zolar soni chegarasi", bool(xato) and "a'zo" in (xato or ""),
          f"xato={xato}")

    # --- ICHMA-ICH ARXIV ----------------------------------------------
    ichki = _zip_yasa([("ichki.txt", "ICHKIMATN " * 30)])
    tashqi = _zip_yasa([("ichki.zip", ichki)])
    matn, _, _, _ = E.extract_zip(tashqi)
    check("1-daraja ichki arxiv ochiladi", "ICHKIMATN" in matn, matn[:80])
    matn, _, _, _ = E.extract_zip(_zip_yasa([("o2.zip", tashqi)]))
    check("2-daraja ichki arxiv OCHILMAYDI", "ICHKIMATN" not in matn,
          matn[:80])

    # --- NOSOZ ARXIV --------------------------------------------------
    _, _, _, xato = E.extract_zip(b"PK\x03\x04buzilgan")
    check("buzilgan arxiv: xato qaytadi, yiqilmaydi", bool(xato), f"{xato}")
    _, _, _, xato = E.extract_zip(
        _zip_yasa([("a.png", b"\x89PNG"), ("b.exe", b"MZ")]))
    check("o'qiladigan hujjat yo'q -> aniq sabab",
          bool(xato) and "o'qiladigan" in (xato or ""), f"{xato}")

    # --- MATN CHIQADI -------------------------------------------------
    matn, _, chiqargich, xato = E.extract_zip(
        _zip_yasa([("shartnoma.txt", "Kafolat muddati 12 oy. " * 5)]))
    check("oddiy a'zodan matn olinadi", "Kafolat muddati" in matn, f"{xato}")
    check("a'zo nomi sarlavha bo'lib qo'shiladi", "shartnoma.txt" in matn)
    check("chiqargich a'zo sonini ko'rsatadi",
          "zipfile(1/1)" == chiqargich, str(chiqargich))

    # --- SEHRLI BAYTLAR -----------------------------------------------
    section("3c. Sehrli baytlar — kengaytma yolg'on bo'lishi mumkin")
    # O'LCHANGAN (2026-08-25): `.doc` deb belgilangan fayllarning bir
    # qismi aslida `docx`. Kengaytmaga ishonsak, o'qiladigan hujjatni
    # bekorga rad etardik.
    check("pdf", E.sniff_magic(b"%PDF-1.4 xxx", "doc") == "pdf")
    check("ole2 (eski .doc)",
          E.sniff_magic(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "doc") == "ole2")
    check("rar", E.sniff_magic(b"Rar!\x1a\x07\x01\x00", "zip") == "rar")
    check("oddiy zip", E.sniff_magic(_zip_yasa([("a.txt", "x")]), "doc") == "zip")
    check("noma'lum bayt -> kengaytma qoladi",
          E.sniff_magic(b"\x00\x01\x02\x03", "pdf") == "pdf")
    check("ole2 qo'llab-quvvatlanadi", E.is_supported("ole2"))
    check("zip qo'llab-quvvatlanadi", E.is_supported("zip"))
    check("rar hamon qo'llab-quvvatlanmaydi", not E.is_supported("rar"))


def test_doc_ole2() -> None:
    """Eski .doc (OLE2) — kutubxonasiz ajratgich.

    Word 97 matnni UTF-16LE bo'lib saqlaydi. Sun'iy fayl quramiz:
    OLE2 imzosi + binar shovqin + haqiqiy matn.
    """
    section("3d. Eski .doc (OLE2) — kutubxonasiz ajratish")

    MATN = ("XIZMAT KO'RSATISH SHARTNOMASI. Kafolat muddati 12 oy "
            "qilib belgilanadi va tovar qabul qilingan kundan boshlanadi. "
            "To'lov shartlari: 30 kun ichida amalga oshiriladi.")
    ole = (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"      # OLE2 imzosi
           + b"\x00\x01\x02\x03" * 40                  # binar shovqin
           + MATN.encode("utf-16-le")
           + b"\xff\xfe" * 30)

    check("sniff_magic OLE2 ni taniydi", E.sniff_magic(ole, "doc") == "ole2")
    matn, _, chiqargich, xato = E.extract_doc(ole)
    check("shartnoma matni ajratildi", "SHARTNOMASI" in matn,
          f"xato={xato} {matn[:70]}")
    check("kafolat bandi to'liq chiqdi", "Kafolat muddati 12 oy" in matn,
          matn[:90])
    check("chiqargich nomi", chiqargich == "ole2-xom", str(chiqargich))

    # --- SHOVQIN FILTRI ---
    # O'LCHANGAN: namunalarning birida matn `яяяяяя...` bilan
    # boshlanardi — Word ning ichki binar maydoni cp1251 da shunday
    # o'qilgan. Uzun, lekin ma'nosiz.
    check("bir xil belgi takrori shovqin deb sanaladi",
          E._ole_shovqinmi("я" * 80))
    check("bo'sh bo'lak shovqin", E._ole_shovqinmi("   "))
    check("haqiqiy matn shovqin EMAS",
          not E._ole_shovqinmi("Kafolat muddati 12 oy bo'ladi va "
                               "to'lov 30 kun ichida"))

    shovqinli = (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
                 + ("я" * 200).encode("utf-16-le")
                 + MATN.encode("utf-16-le"))
    matn, _, _, _ = E.extract_doc(shovqinli)
    check("shovqin natijadan chiqarildi", "я" * 40 not in matn, matn[:70])
    check("shovqin yonidagi matn saqlandi", "Kafolat muddati" in matn,
          matn[:70])

    # --- MATNSIZ FAYL ---
    _, _, _, xato = E.extract_doc(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"
                                  + b"\x00" * 500)
    check("matnsiz OLE2 -> aniq sabab", bool(xato), str(xato))


def test_status_logic() -> None:
    section("3. Format / status mantig'i (tarmoqsiz)")

    check("rar -> unsupported", not E.is_supported("rar"))
    check("doc (eski binar) -> unsupported", not E.is_supported("doc"))
    check("xls (eski binar) -> unsupported", not E.is_supported("xls"))
    check("pdf/docx/xlsx/html -> supported",
          all(E.is_supported(x) for x in ("pdf", "docx", "xlsx", "html")))

    check("kengaytma nomdan olindi",
          E.sniff_ext({"file_type": None, "name": "TZ.PDF"}) == "pdf")
    check("kengaytma content_type dan olindi",
          E.sniff_ext({"file_type": "", "name": "hujjat",
                       "content_type": "application/pdf"}) == "pdf")

    # unsupported — TARMOQQA CHIQMASDAN hal bo'lishi kerak
    rec = E.process(None, {"tender_id": 1, "file_ref": "x", "file_type": "rar",
                           "name": "arxiv.rar", "size_bytes": 1000,
                           "source_platform": "xt-xarid", "file_id": "abc"})
    check("unsupported yuklab olmasdan aniqlandi", rec["status"] == "unsupported",
          rec["error"] or "")

    # too_large — metama'lumot bo'yicha, yuklab olmasdan
    rec = E.process(None, {"tender_id": 1, "file_ref": "x", "file_type": "pdf",
                           "name": "katta.pdf", "size_bytes": E.MAX_BYTES + 1,
                           "source_platform": "xt-xarid", "file_id": "abc"})
    check("too_large yuklab olmasdan aniqlandi", rec["status"] == "too_large",
          rec["error"] or "")

    # CHIZMA/SKAN PDF taqlidi: belgi ko'p, HARF yo'q -> unreadable.
    # `download` ni vaqtincha almashtiramiz (tarmoqqa chiqilmaydi).
    real_download = E.download
    row = {"tender_id": 1, "file_ref": "x", "file_type": "txt", "name": "a.txt",
           "size_bytes": 100, "source_platform": "xt-xarid", "file_id": "abc"}
    try:
        E.download = lambda s, r: (b"No No No 296 296 83 88937 123 456 789 000", None)
        rec = E.process(None, row)
        check("harfsiz 'matn' -> unreadable", rec["status"] == "unreadable",
              (rec["error"] or "")[:60])

        E.download = lambda s, r: (("Texnik topshiriq: kompressor yetkazib berish, "
                                    "kafolat muddati yigirma turt oy").encode(), None)
        rec = E.process(None, row)
        check("haqiqiy matn -> ok", rec["status"] == "ok", f"{rec['char_count']} belgi")

        E.download = lambda s, r: (b"", None)
        rec = E.process(None, row)
        check("bo'sh fayl -> unreadable", rec["status"] == "unreadable",
              (rec["error"] or "")[:40])
    finally:
        E.download = real_download

    # yuklab olish manzili yo'q -> download_failed
    rec = E.process(None, {"tender_id": 1, "file_ref": "x", "file_type": "pdf",
                           "name": "a.pdf", "size_bytes": 100,
                           "source_platform": "noma'lum", "file_id": None,
                           "file_path": None})
    check("manzilsiz hujjat -> download_failed", rec["status"] == "download_failed",
          rec["error"] or "")


# ---------------------------------------------------------------------------
# 4. HAQIQIY HUJJATLAR (tarmoq) — har turdan bir nechta
# ---------------------------------------------------------------------------
SAMPLE_SQL = """
SELECT d.tender_id, d.file_ref, d.file_id, d.file_path, d.name,
       d.size_bytes, d.content_type, d.file_type, d.source_platform
FROM tender_document d
WHERE lower(d.file_type) = %(ft)s AND coalesce(d.size_bytes, 0) < 8000000
ORDER BY d.tender_id DESC
LIMIT %(n)s
"""


def test_live(conn, session) -> None:
    section("4. Haqiqiy hujjatlar (manbadan yuklab olinadi)")
    for ft, n in (("pdf", 3), ("docx", 3), ("xlsx", 2)):
        with conn.cursor(cursor_factory=RealDictCursor) as cur:
            cur.execute(SAMPLE_SQL, {"ft": ft, "n": n})
            rows = [dict(r) for r in cur.fetchall()]
        if not rows:
            check(f"{ft}: namuna topilmadi", False)
            continue
        recs = [E.process(session, r) for r in rows]
        ok = [r for r in recs if r["status"] == "ok"]
        chars = sum(r["char_count"] or 0 for r in ok)
        check(f"{ft.upper()}: {len(rows)} tadan {len(ok)} tasidan matn chiqdi",
              len(ok) >= 1, f"{chars:,} belgi")
        for r in recs:
            check(f"  {ft} status to'g'ri toifada",
                  r["status"] in ("ok", "unreadable", "unsupported",
                                  "too_large", "download_failed"), r["status"])
            if r["status"] == "ok":
                check("  ok bo'lsa matn bor va char_count mos",
                      r["text"] and r["char_count"] == len(r["text"]))
            else:
                check("  ok bo'lmasa sabab yozilgan (qo'lda tekshirish uchun)",
                      bool(r["error"]), (r["error"] or "")[:50])
            break   # har turdan bitta batafsil tekshiruv yetarli


# ---------------------------------------------------------------------------
# 5. KESH — qayta yurgizishda takroriy yuklab olish bo'lmasligi
# ---------------------------------------------------------------------------
class _Args:
    def __init__(self, **kw):
        self.limit = self.tender_id = self.platform = self.file_type = None
        self.force = False
        for k, v in kw.items():
            setattr(self, k, v)


#: Sinov yaratadigan qatorlar shu prefiks bilan boshlanadi —
#: haqiqiy hujjatlardan ajratish va tozalash uchun.
SOXTA_PREFIKS = "[SINOV]/doctext/"


def test_navbat_adolati(conn) -> None:
    """NAVBAT TARTIBI bir platformani ochlikka mahkum qilmasin.

    O'LCHANGAN NUQSON (2026-09-03). `fetch_targets()` da tartib
    `ORDER BY d.tender_id DESC` edi. Ikki platformaning ID fazolari
    esa kesishmaydi:

        xt-xarid          108 .. 8 538 264
        uzex       20 000 475 229 .. 20 000 510 026

    Ya'ni HAR uzex hujjati HAR xt-xarid hujjatidan oldin turardi.
    O'lchangan navbat o'rinlari: uzex 1..342, xt-xarid 343..906.
    Qadamda vaqt byudjeti bor (~25 daqiqa to'liq qamrov), shuning
    uchun xt-xarid hujjatlariga NAVBAT YETIB BORMASDI —
    matn qamrovi 31/595, uzex esa 2 925 ta.

    Bu sinov FILTRNI emas, TARTIBNI tekshiradi: filtr ham, byudjet
    ham to'g'ri edi.
    """
    section("5a. Navbat adolati — katta ID li platforma ochlik qilmasin")

    # Ikki soxta platforma, ID fazolari AYNAN haqiqiydek uzoq.
    KATTA = [(20_000_900_001 + i, "zzbig") for i in range(3)]
    KICHIK = [(900_001 + i, "zzsmall") for i in range(3)]
    yaratilgan = []
    try:
        with conn.cursor() as cur:
            for tid, plat in KATTA + KICHIK:
                cur.execute(
                    "INSERT INTO tender (id, source_id, source_platform, "
                    "  status, close_at, raw_json) "
                    "VALUES (%s, %s, %s, 'open', "
                    "        now() + interval '7 days', '{}') "
                    "ON CONFLICT (id) DO UPDATE SET status='open', "
                    "  close_at = now() + interval '7 days'",
                    (tid, tid, plat))
                ref = f"{SOXTA_PREFIKS}navbat-{tid}.pdf"
                cur.execute(
                    "INSERT INTO tender_document (tender_id, file_ref, "
                    "  source_platform, fetched_at, holat, file_type, urinish) "
                    "VALUES (%s, %s, %s, now(), 'navbatda', 'pdf', 0) "
                    "ON CONFLICT (tender_id, file_ref) DO NOTHING",
                    (tid, ref, plat))
                yaratilgan.append(tid)
            conn.commit()

        hammasi = E.fetch_targets(conn, _Args())
        soxta = [r for r in hammasi
                 if r["source_platform"] in ("zzbig", "zzsmall")]
        check("soxta qatorlar navbatda ko'rindi", len(soxta) == 6,
              f"{len(soxta)} ta")
        if len(soxta) != 6:
            return

        platformalar = [r["source_platform"] for r in soxta]
        # ASOSIY TEKSHIRUV: kichik ID li platforma birinchi UCHTALIKKA
        # tushsin. Eski tartibda u 4-o'rindan oldin CHIQMASDI.
        birinchi_kichik = platformalar.index("zzsmall")
        check("kichik ID li platforma OXIRGA surilmadi",
              birinchi_kichik <= 2,
              f"birinchi 'zzsmall' {birinchi_kichik + 1}-o'rinda "
              f"(eski tartibda 4-o'rin edi): {platformalar}")

        # Platforma ICHIDA tartib saqlansin — yangi tender oldin.
        kichik_idlar = [r["tender_id"] for r in soxta
                        if r["source_platform"] == "zzsmall"]
        check("platforma ICHIDA tartib saqlandi (yangi tender oldin)",
              kichik_idlar == sorted(kichik_idlar, reverse=True),
              str(kichik_idlar))

        # Tartib DETERMINISTIK bo'lsin — checkpoint/qayta yurish shunga
        # tayanadi. Ikki chaqiruv bir xil ketma-ketlik berishi shart.
        yana = [r["file_ref"] for r in E.fetch_targets(conn, _Args())
                if r["source_platform"] in ("zzbig", "zzsmall")]
        check("tartib DETERMINISTIK (ikki chaqiruv bir xil)",
              yana == [r["file_ref"] for r in soxta])
    finally:
        with conn.cursor() as cur:
            cur.execute("DELETE FROM tender_document WHERE file_ref LIKE %s",
                        (SOXTA_PREFIKS + "navbat-%",))
            cur.execute("DELETE FROM tender WHERE source_platform "
                        "IN ('zzbig','zzsmall')")
        conn.commit()


def test_cache(conn) -> None:
    section("5. Kesh (takroriy yuklab olmaslik)")
    with conn.cursor() as cur:
        cur.execute("SELECT count(*) FROM tender_document_text")
        done = cur.fetchone()[0]
    if done == 0:
        check("bazada ishlangan hujjat bor", False, "avval ETL ni yurgizing")
        return

    # Ishlangan hujjat --force siz QAYTA tanlanmasligi kerak
    with conn.cursor() as cur:
        cur.execute("SELECT tender_id FROM tender_document_text LIMIT 1")
        tid = cur.fetchone()[0]

    without = E.fetch_targets(conn, _Args(tender_id=tid))
    with_force = E.fetch_targets(conn, _Args(tender_id=tid, force=True))
    check("ishlangan hujjat --force siz tanlanmaydi", len(without) < len(with_force),
          f"#{tid}: --force siz {len(without)}, --force bilan {len(with_force)}")

    processed_refs = {r["file_ref"] for r in with_force} - {r["file_ref"] for r in without}
    check("tanlanmaganlar aynan bazadagilar", len(processed_refs) > 0,
          f"{len(processed_refs)} ta o'tkazib yuborildi")

    # Idempotentlik: bir yozuvni ikki marta saqlash dublikat
    # yaratmasligi kerak.
    #
    # SOXTA QATORDA ishlaydi. ILGARI BU YERDA HAQIQIY hujjat
    # olinardi va u IKKI ZARAR berardi (o'lchangan 2026-09-01, M-1):
    #
    #   1. `E.save(rec)` haqiqiy ajratilgan matnni `"sinov"` (5
    #      belgi) bilan ALMASHTIRARDI — ma'lumot yo'qolardi.
    #   2. Keyin matn qatori o'chirilardi, lekin
    #      `tender_document.holat` `ok` bo'lib QOLARDI. Hujjat
    #      `ok` bo'lgani uchun ETL uni QAYTA OLMASDI
    #      (`fetch_targets()` `ok` larni o'tkazadi) — ya'ni matn
    #      butunlay yo'qolardi.
    #
    #   sinovdan OLDIN 29 ta dalilsiz `ok`, KEYIN 30 ta (+1 har
    #   yurishda).
    #
    # Endi baza ham buni to'xtatadi (`hujjat_ok_dalil_trg`), lekin
    # sinov HAQIQIY ma'lumotga UMUMAN tegmasligi kerak.
    soxta_ref = f"{SOXTA_PREFIKS}kesh-idempotent.pdf"
    with conn.cursor() as cur:
        cur.execute("""INSERT INTO tender_document
                         (tender_id, file_ref, name, file_type,
                          source_platform, holat)
                       VALUES (%s, %s, '[SINOV] kesh', 'pdf',
                               'xt-xarid', 'navbatda')
                       ON CONFLICT (tender_id, file_ref) DO NOTHING""",
                    (tid, soxta_ref))
    conn.commit()
    rec = {"tender_id": tid, "file_ref": soxta_ref,
           "text": "sinov matni", "status": "ok", "char_count": 11,
           "page_count": 1, "error": None, "extractor": "plain"}
    try:
        E.save(conn, rec)
        E.save(conn, rec)
        with conn.cursor() as cur:
            cur.execute("""SELECT count(*) FROM tender_document_text
                           WHERE tender_id = %s AND file_ref = %s""",
                        (tid, soxta_ref))
            cnt = cur.fetchone()[0]
        check("ikki marta saqlash dublikat yaratmadi", cnt == 1, f"qator: {cnt}")
        with conn.cursor() as cur:
            cur.execute("""SELECT holat FROM tender_document
                           WHERE tender_id = %s AND file_ref = %s""",
                        (tid, soxta_ref))
            check("metadata holati ham yangilandi",
                  cur.fetchone()[0] == "ok")
    finally:
        # IKKALA jadval ham tozalanadi. Faqat matnni o'chirish
        # `holat='ok'` ni DALILSIZ qoldirardi — aynan tuzatilgan
        # nuqson.
        with conn.cursor() as cur:
            cur.execute("""DELETE FROM tender_document_text
                           WHERE tender_id = %s AND file_ref = %s""",
                        (tid, soxta_ref))
            cur.execute("""DELETE FROM tender_document
                           WHERE tender_id = %s AND file_ref = %s""",
                        (tid, soxta_ref))
        conn.commit()

    # QO'ROVUL O'ZI ISHLAYDIMI: dalilsiz `ok` RAD ETILSIN.
    # Bu "hech narsa buzilmadi" bilan "qo'rovul bor" ni ajratadi.
    import psycopg2
    with conn.cursor() as cur:
        cur.execute("""INSERT INTO tender_document
                         (tender_id, file_ref, name, file_type,
                          source_platform, holat)
                       VALUES (%s, %s, '[SINOV] qorovul', 'pdf',
                               'xt-xarid', 'navbatda')""",
                    (tid, SOXTA_PREFIKS + "qorovul.pdf"))
    conn.commit()
    try:
        with conn.cursor() as cur:
            cur.execute("""UPDATE tender_document SET holat='ok'
                           WHERE tender_id=%s AND file_ref=%s""",
                        (tid, SOXTA_PREFIKS + "qorovul.pdf"))
        conn.commit()
        check("dalilsiz `ok` BAZA darajasida rad etiladi", False,
              "qabul qilindi!")
    except psycopg2.Error as e:
        conn.rollback()
        check("dalilsiz `ok` BAZA darajasida rad etiladi", True,
              str(e).splitlines()[0][:80])
    finally:
        with conn.cursor() as cur:
            cur.execute("""DELETE FROM tender_document
                           WHERE tender_id=%s AND file_ref=%s""",
                        (tid, SOXTA_PREFIKS + "qorovul.pdf"))
        conn.commit()

    # NOMUVOFIQLIK O'LCHOVI — qoldiq qolmasin.
    with conn.cursor() as cur:
        cur.execute("SELECT ok_dalilsiz, ok_status_qarama_qarshi "
                    "FROM v_hujjat_dalil_nomuvofiq")
        dalilsiz, qarama = cur.fetchone()
    check("dalilsiz `ok` hujjat YO'Q", dalilsiz == 0, f"{dalilsiz} ta")
    check("`ok` yorlig'i matn statusiga ZID emas", qarama == 0, f"{qarama} ta")


# ---------------------------------------------------------------------------
# 6. STATUS TAQSIMOTI — bazadagi haqiqiy holat
# ---------------------------------------------------------------------------
def report(conn) -> None:
    section("6. Bazadagi status taqsimoti")
    with conn.cursor() as cur:
        cur.execute("""SELECT status, count(*), coalesce(sum(char_count), 0)
                       FROM tender_document_text GROUP BY status ORDER BY 2 DESC""")
        rows = cur.fetchall()
        cur.execute("SELECT count(*) FROM tender_document")
        total_docs = cur.fetchone()[0]

    total = sum(r[1] for r in rows)
    print(f"  Jami hujjat metama'lumoti : {total_docs}")
    print(f"  Matn ajratish o'tkazilgan : {total} ({total * 100 // max(total_docs, 1)}%)")
    print(f"  {'status':<16}{'soni':>7}{'belgi':>14}")
    for st, n, ch in rows:
        print(f"  {st:<16}{n:>7}{ch:>14,}")
    manual = sum(n for st, n, _ in rows if st != "ok")
    ok_n = sum(n for st, n, _ in rows if st == "ok")
    print(f"\n  -> o'qildi                : {ok_n}")
    print(f"  -> qo'lda tekshirish kerak: {manual}")

    check("kamida bitta hujjat 'ok'", ok_n > 0)
    check("har bir yozuvda status bor", total == sum(r[1] for r in rows))
    with conn.cursor() as cur:
        cur.execute("""SELECT count(*) FROM tender_document_text
                       WHERE status = 'ok' AND (text IS NULL OR char_count IS NULL)""")
        bad = cur.fetchone()[0]
    check("'ok' yozuvlarda matn bo'sh emas", bad == 0, f"nosoz: {bad}")
    with conn.cursor() as cur:
        cur.execute("""SELECT count(*) FROM tender_document_text
                       WHERE status <> 'ok' AND error IS NULL""")
        bad = cur.fetchone()[0]
    check("'ok' bo'lmaganlarda sabab ko'rsatilgan", bad == 0, f"sababsiz: {bad}")


def main() -> None:
    ap = argparse.ArgumentParser(description="Hujjat matni sinovi")
    rejim.bayroqlar(ap)
    args = rejim.moslash(ap.parse_args())

    dsn = os.environ.get("XT_DB_DSN")
    if not dsn:
        sys.exit("XATO: XT_DB_DSN yo'q (.env ni tekshiring).")

    conn = psycopg2.connect(dsn)
    try:
        test_schema(conn)
        test_parsers()
        test_zip()
        test_doc_ole2()
        test_status_logic()
        if not args.tarmoqsiz:
            import requests
            test_live(conn, requests.Session())
        else:
            print("\n=== 4. Haqiqiy hujjatlar — O'TKAZILDI (--offline) ===")
        test_navbat_adolati(conn)
        test_cache(conn)
        report(conn)
    finally:
        conn.close()

    print(f"\n{'=' * 46}\nNATIJA: {PASS} o'tdi, {FAIL} yiqildi")
    sys.exit(1 if FAIL else 0)


if __name__ == "__main__":
    main()
