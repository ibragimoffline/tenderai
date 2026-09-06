#!/usr/bin/env python3
"""
HUJJAT MATNI ETL  (P0-2)
========================
Tenderga biriktirilgan fayllarni MANBADAN yuklab oladi, matnini DETERMINISTIK
parserlar bilan ajratib oladi va `tender_document_text` ga yozadi.

NEGA KERAK: `tender_document` da faqat metama'lumot bor (nom, hajm, tur).
Tenderning haqiqiy mazmuni — texnik topshiriq, talablar, shartlar — PDF/DOCX
ichida qulflangan. Matn ajratilmaguncha na qidiruv, na talab ajratish ishlaydi.

MUHIM: bu skript AI/model CHAQIRMAYDI. Faqat parserlar:
    pdf  -> pypdf            docx -> python-docx
    xlsx -> openpyxl         txt/csv/html/xml -> stdlib
Qolgan formatlar (rar, zip, doc, ...) -> 'unsupported'.

FAYLNI QAYERDAN OLAMIZ (mantiq `api/main.py` dan nusxa olingan — u fayl
o'zgartirilmagan):
    xt-xarid : GET  https://api.xt-xarid.uz/file/<file_id>
    uzex     : POST https://apietender.uzex.uz/api/common/DownloadFile?path=...
               (GET -> 405, brauzer User-Agent shart)
Fayl DISKKA saqlanmaydi — oqim bilan olinadi, matni ajratiladi, o'zi tashlanadi.

STATUSLAR (TZ talabi: "o'qib bo'lmaydigan fayllar 'qo'lda tekshirish talab
etiladi' deb belgilanadi" — 'ok' dan boshqa hammasi shu toifaga kiradi):
    ok | unreadable | unsupported | download_failed | too_large

Ishga tushirish:
    python etl_doc_text.py                      # ishlanmagan barcha hujjatlar
    python etl_doc_text.py --limit 20           # sinov uchun
    python etl_doc_text.py --tender-id 1493304  # bitta tender
    python etl_doc_text.py --force              # qayta ishlash (kesh bekor)
    python etl_doc_text.py --limit 5 --dry-run  # DBga yozmaydi
"""
import argparse
import csv
import io
import logging
import os
import re
import sys
import time
import zipfile
from html.parser import HTMLParser
from typing import Any, Dict, List, Optional, Tuple

import zipfile

import requests

import etl_ishonch as ish

try:
    from dotenv import load_dotenv
except ImportError:  # dotenv bo'lmasa ham muhit o'zgaruvchisidan ishlaydi
    load_dotenv = None

try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
except ImportError:
    psycopg2 = None

# --- Parserlar. Yo'q bo'lsa skript ishlayveradi, o'sha format 'unsupported' ---
try:
    from pypdf import PdfReader
    # pypdf nostandart PDF haqida HAR SAHIFADA ogohlantirish yozadi va
    # natija satrlarini bosib ketadi. Muammo baribir `status` da aks etadi.
    logging.getLogger("pypdf").setLevel(logging.ERROR)
except ImportError:
    PdfReader = None
try:
    import docx as python_docx
except ImportError:
    python_docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None


# ---------------------------------------------------------------------------
# Sozlamalar
# ---------------------------------------------------------------------------
REQUEST_DELAY = 0.6          # manbaga bosim qilmaslik uchun (soniya)
TIMEOUT       = 90           # yuklab olish (katta PDF'lar bor)
MAX_RETRIES   = 3
RETRY_BACKOFF = 2.0
CHUNK         = 65536        # oqim bo'lagi

MAX_BYTES     = 25 * 1024 * 1024    # 25 MB dan katta -> too_large
MAX_CHARS     = 400_000             # bazaga yoziladigan matn chegarasi

#: VAQT BYUDJETI (soniya). `run_etl.py` uni `--max-seconds` bilan
#: beradi; yakka yurgizilganda standart shu. 0 = cheksiz.
#:
#: `etl_uzex.py` va `etl_details.py` da byudjet ALLAQACHON bor edi,
#: bu qadamda esa YO'Q edi — holbuki u quvurdagi ENG SEKINI
#: (to'liq qamrov ~25 daqiqa). Natija: host CTRL+C yuborganda
#: (`0xC000013A`) qadam fayl o'rtasida o'lardi.
STANDART_BYUDJET = 1200

#: Byudjet/to'xtash so'rovi bilan TUGALLANMAGAN yurish. `0` (ok) ham,
#: `1` (xato) ham EMAS — `run_etl.py` uni `partial` deb o'qiydi.
CHIQISH_QISMAN = 7

#: To'xtash so'rovini yig'uvchi (signal + byudjet). `main()` da quriladi.
_TOXTATGICH: Optional["ish.Toxtatgich"] = None
MIN_CHARS     = 20                  # shundan kam matn = amalda bo'sh (skan)
# Chizma/skan PDF'lardan ba'zan bir necha belgi "sizib" chiqadi:
#   '№ № № ³ Ø296 Ø296Ø83Ø88937'  — uzunligi MIN_CHARS dan katta, lekin bu
# MATN emas. Shuning uchun HARF sonini alohida talab qilamiz.
MIN_LETTERS   = 40
PDF_MAX_PAGES = 300                 # 1000 sahifali kitoblar bor — vaqtni cheklaymiz
XLSX_MAX_ROWS = 5000                # varaqdagi qatorlar chegarasi

# `api/main.py` dagi bilan bir xil (o'sha fayl o'zgartirilmadi)
_FILE_URL      = {"xt-xarid": "https://api.xt-xarid.uz/file/{file_id}"}
_POST_DOWNLOAD = {"uzex": "https://apietender.uzex.uz/api/common/DownloadFile"}
_BROWSER_UA    = ("Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
                  "AppleWebKit/537.36 (KHTML, like Gecko) "
                  "Chrome/120.0.0.0 Safari/537.36")

# Qaysi kengaytmani qaysi parser oladi
EXT_PDF   = {"pdf"}
EXT_DOCX  = {"docx", "docm"}
EXT_XLSX  = {"xlsx", "xlsm"}
EXT_PLAIN = {"txt", "csv", "htm", "html", "xml", "json", "md"}
EXT_ZIP   = {"zip"}
EXT_OLE2  = {"ole2"}      # eski Word/Excel — `sniff_magic` beradi

# --- ZIP CHEKLOVLARI -------------------------------------------------
# Arxiv TASHQI MANBADAN keladi, ya'ni ishonchsiz.
#
# ZIP SLIP (a'zo nomida `../../`) bu yerda TUZILISHIGA KO'RA mumkin
# emas: biz hech qachon diskka yozmaymiz, a'zo faqat xotirada
# o'qiladi va nomi faqat YORLIQ sifatida ishlatiladi. Shunday bo'lsa
# ham nom tozalanadi — chalg'ituvchi matn chiqmasin.
#
# ZIP BOMBA esa haqiqiy xavf: 40 KB arxiv ochilganda gigabaytga
# aylanishi mumkin. Uchta chegara qo'yamiz va ularning HAMMASI
# ochishdan OLDIN, `ZipInfo.file_size` (siqilmagan hajm) bo'yicha
# tekshiriladi.
ZIP_MAX_MEMBERS = 40          # a'zolar soni
ZIP_MAX_TOTAL   = 60 * 1024 * 1024   # jami siqilmagan hajm
ZIP_MAX_MEMBER  = 25 * 1024 * 1024   # bitta a'zo (MAX_BYTES bilan bir xil)
# Ichma-ich arxiv OCHILMAYDI: 1 daraja yetarli, chuqurroq ketish
# bombaga eshik ochadi va amalda uchramaydi.
ZIP_MAX_DEPTH   = 1

#: Sehrli baytlar — KENGAYTMA YOLG'ON BO'LISHI MUMKIN.
#: Namunada (2026-08-25) `.doc` deb belgilangan 6 fayldan 1 tasi
#: aslida `docx` (ZIP) bo'lib chiqdi. Kengaytmaga ishonsak, o'qilishi
#: mumkin bo'lgan hujjatni bekorga rad etardik.
_MAGIC = [
    (b"%PDF", "pdf"),
    (b"PK\x03\x04", "zip"),           # zip | docx | xlsx — pastda ajratiladi
    (b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1", "ole2"),   # eski doc/xls
    (b"Rar!\x1a\x07", "rar"),
    (b"7z\xbc\xaf\x27\x1c", "7z"),
    (b"{\\rtf", "rtf"),
    (b"\x1f\x8b", "gz"),
]

# Bilib turib rad etamiz (arxiv / eski binar format / rasm)
EXT_KNOWN_UNSUPPORTED = {"rar", "7z", "tar", "gz", "xls", "ppt",
                         "pptx", "jpg", "jpeg", "png", "gif", "tif", "tiff",
                         "bmp", "exe", "dwg", "sig", "p7s", "pfx", "ofd"}

_WS_RE = re.compile(r"[ \t ]+")
_NL_RE = re.compile(r"\n{3,}")


# ---------------------------------------------------------------------------
# Yuklab olish
# ---------------------------------------------------------------------------
def download(session: requests.Session, row: dict) -> Tuple[Optional[bytes], Optional[str]]:
    """Faylni manbadan oladi. -> (baytlar, xato).

    Oqim bilan o'qiladi va MAX_BYTES da UZILADI — server Content-Length
    bermasa ham katta fayl xotirani to'ldirmaydi.
    """
    platform = row.get("source_platform") or "xt-xarid"

    if platform == "xt-xarid" and row.get("file_id"):
        method, url, params = "GET", _FILE_URL["xt-xarid"].format(file_id=row["file_id"]), None
    elif platform in _POST_DOWNLOAD and row.get("file_path"):
        method, url, params = "POST", _POST_DOWNLOAD[platform], {"path": row["file_path"]}
    else:
        return None, f"'{platform}' uchun yuklab olish manzili yo'q"

    for attempt in range(1, MAX_RETRIES + 1):
        try:
            resp = session.request(method, url, params=params,
                                   headers={"User-Agent": _BROWSER_UA},
                                   stream=True, timeout=TIMEOUT)
            resp.raise_for_status()

            # Sarlavhada hajm bo'lsa — yuklamasdan turib rad etamiz
            clen = resp.headers.get("Content-Length")
            if clen and clen.isdigit() and int(clen) > MAX_BYTES:
                resp.close()
                return None, f"__TOO_LARGE__{clen}"

            buf = io.BytesIO()
            size = 0
            for part in resp.iter_content(chunk_size=CHUNK):
                if not part:
                    continue
                size += len(part)
                if size > MAX_BYTES:
                    resp.close()
                    return None, f"__TOO_LARGE__{size}+"
                buf.write(part)
            resp.close()
            return buf.getvalue(), None
        except requests.RequestException as e:
            if attempt == MAX_RETRIES:
                return None, str(e)[:400]
            time.sleep(RETRY_BACKOFF ** attempt)
    return None, "noma'lum xato"


# ---------------------------------------------------------------------------
# Matn ajratish — har biri (matn, sahifa_soni, extractor, xato) qaytaradi
# ---------------------------------------------------------------------------
def clean(text: str) -> str:
    """Ortiqcha bo'shliq/qator — bazada joy va keyingi tahlilda shovqin."""
    # NUL bayt — PostgreSQL TEXT uni QABUL QILMAYDI (ValueError bilan yiqiladi).
    # Buzilgan PDF/DOCX dan chiqib qolishi mumkin, shuning uchun birinchi olib
    # tashlaymiz. Boshqa boshqaruv belgilari ham matnda keraksiz.
    text = text.replace("\x00", " ")
    text = "".join(ch if ch >= " " or ch in "\n\r\t" else " " for ch in text)
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = _WS_RE.sub(" ", text)
    text = "\n".join(line.strip() for line in text.split("\n"))
    return _NL_RE.sub("\n\n", text).strip()


def extract_pdf(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if PdfReader is None:
        return "", None, "pypdf", "pypdf o'rnatilmagan"
    try:
        reader = PdfReader(io.BytesIO(data))
        if reader.is_encrypted:
            # Bo'sh parol ko'p hollarda ochadi (himoya faqat tahrirdan)
            try:
                reader.decrypt("")
            except Exception:  # noqa: BLE001
                return "", None, "pypdf", "PDF parol bilan himoyalangan"
        pages = reader.pages
        total = len(pages)
        parts: List[str] = []
        for i, page in enumerate(pages):
            if i >= PDF_MAX_PAGES:
                break
            try:
                parts.append(page.extract_text() or "")
            except Exception:  # noqa: BLE001
                continue    # bitta sahifa buzilgan bo'lsa qolganini olamiz
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
        return clean("\n".join(parts)), total, "pypdf", None
    except Exception as e:  # noqa: BLE001
        return "", None, "pypdf", str(e)[:400]


def extract_docx(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if python_docx is None:
        return "", None, "python-docx", "python-docx o'rnatilmagan"
    try:
        doc = python_docx.Document(io.BytesIO(data))
        parts = [p.text for p in doc.paragraphs if p.text and p.text.strip()]
        n_para = len(doc.paragraphs)
        # JADVALLAR — texnik topshiriqda talablar ko'pincha aynan jadvalda
        for table in doc.tables:
            for r in table.rows:
                cells = [c.text.strip() for c in r.cells]
                line = " | ".join(x for x in cells if x)
                if line:
                    parts.append(line)
        return clean("\n".join(parts)), n_para, "python-docx", None
    except (zipfile.BadZipFile, KeyError) as e:
        # .docx nomi bilan aslida .doc/.rtf bo'lishi mumkin
        return "", None, "python-docx", f"DOCX emas yoki buzilgan: {e}"[:400]
    except Exception as e:  # noqa: BLE001
        return "", None, "python-docx", str(e)[:400]


def extract_xlsx(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    if openpyxl is None:
        return "", None, "openpyxl", "openpyxl o'rnatilmagan"
    wb = None
    try:
        # read_only + values_only — formula/uslub o'qilmaydi, xotira tejaladi
        wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
        parts: List[str] = []
        for ws in wb.worksheets:
            parts.append(f"[{ws.title}]")
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= XLSX_MAX_ROWS:
                    break
                cells = [str(c).strip() for c in row if c is not None and str(c).strip()]
                if cells:
                    parts.append(" | ".join(cells))
            if sum(len(p) for p in parts) > MAX_CHARS:
                break
        return clean("\n".join(parts)), len(wb.worksheets), "openpyxl", None
    except (zipfile.BadZipFile, KeyError) as e:
        return "", None, "openpyxl", f"XLSX emas yoki buzilgan: {e}"[:400]
    except Exception as e:  # noqa: BLE001
        return "", None, "openpyxl", str(e)[:400]
    finally:
        if wb is not None:
            try:
                wb.close()
            except Exception:  # noqa: BLE001
                pass


class _TextHTML(HTMLParser):
    """HTML dan ko'rinadigan matnni oladi (script/style tashlanadi)."""
    _SKIP = {"script", "style", "head", "meta", "link"}

    def __init__(self) -> None:
        super().__init__(convert_charrefs=True)
        self.parts: List[str] = []
        self._skip = 0

    def handle_starttag(self, tag, attrs):
        if tag in self._SKIP:
            self._skip += 1

    def handle_endtag(self, tag):
        if tag in self._SKIP and self._skip:
            self._skip -= 1
        elif tag in ("p", "br", "div", "tr", "li"):
            self.parts.append("\n")

    def handle_data(self, data):
        if not self._skip and data.strip():
            self.parts.append(data)


def extract_plain(data: bytes, ext: str) -> Tuple[str, Optional[int], str, Optional[str]]:
    """txt/csv/html/xml — stdlib. Kodlash noma'lum, ketma-ket sinaymiz."""
    text = None
    for enc in ("utf-8", "utf-8-sig", "cp1251", "cp1252"):
        try:
            text = data.decode(enc)
            break
        except UnicodeDecodeError:
            continue
    if text is None:
        text = data.decode("utf-8", errors="replace")

    if ext in ("htm", "html", "xml"):
        p = _TextHTML()
        try:
            p.feed(text)
            text = " ".join(p.parts)
        except Exception:  # noqa: BLE001
            text = re.sub(r"<[^>]+>", " ", text)   # zaxira: teglarni olib tashlash
    elif ext == "csv":
        try:
            rows = list(csv.reader(io.StringIO(text)))
            text = "\n".join(" | ".join(c.strip() for c in r if c.strip())
                             for r in rows[:XLSX_MAX_ROWS])
        except Exception:  # noqa: BLE001
            pass

    return clean(text), None, "plain", None


#: OLE2 matnida yaroqli deb hisoblanadigan belgilar.
_OLE_YAROQLI = re.compile(
    r"[0-9A-Za-z\u0400-\u04FF\u02BC'\u2018\u2019\u201C\u201D"
    r"\u00AB\u00BB\u2116\s.,;:!?()\[\]/\\%+\-\u2013\u2014=\"#&@*]+")

#: Bo'lak shu uzunlikdan qisqa bo'lsa — tasodifiy shovqin.
OLE_MIN_BOLAK = 40

#: Bir xil belgining uzun takrori — Word ning ichki to'ldiruvchi
#: maydoni. Bo'lakni butunlay rad etish XATO bo'lardi: shovqin
#: HAQIQIY MATNGA YOPISHIB kelishi mumkin va u bilan birga tushib
#: ketardi. Shuning uchun avval KESIB tashlaymiz, keyin baholaymiz.
_TAKROR_RE = re.compile(r"(.)\1{7,}")


def _ole_shovqinmi(bolak: str) -> bool:
    """Bir xil belgi takrori yoki juda kambag'al bo'lak — binar shovqin.

    O'LCHANDI: namunalarning birida matn `яяяяяяяя...` bilan
    boshlanardi — bu Word ning ichki binar maydoni cp1251 da shunday
    o'qilgani. Uzunligi katta, ma'nosi yo'q.
    """
    toza = bolak.strip()
    if not toza:
        return True
    xilma = len(set(toza))
    if xilma <= 3:                       # "яяяя", "0000", "    "
        return True
    # Eng ko'p uchraydigan belgi bo'lakning yarmidan ko'pini egallasa
    eng = max(toza.count(ch) for ch in set(toza))
    return eng > len(toza) * 0.5


def extract_doc(data: bytes) -> Tuple[str, Optional[int], str, Optional[str]]:
    """Eski Word (.doc, OLE2) dan matn — QO'SHIMCHA KUTUBXONASIZ.

    NEGA `olefile` EMAS: Word 97-2003 matnni `WordDocument` oqimida
    UTF-16LE (yoki cp1251) da UZUN UZLUKSIZ bo'laklar sifatida
    saqlaydi. FIB ni to'g'ri o'qish uchun `olefile` kerak, lekin
    RAG uchun bizga matnning O'ZI yetarli — bo'lak chegarasi aniq
    bo'lishi shart emas.

    O'LCHANDI (12 fayl, ochiq tenderlar, 2026-08-25): 11/12 (92%)
    faylda o'qiladigan xarid matni topildi — shartnoma bandlari,
    kafolat va to'lov shartlari. Yangi bog'liqlik uchun bu farq
    yetarli emas edi.

    IKKI KODLASH sinaladi va HARFLAR SONI ko'proq chiqqani olinadi:
    noto'g'ri kodlash harf bermaydi, shuning uchun bu mezon ishonchli.
    """
    eng_matn, eng_ball = "", 0
    for kodlash in ("utf-16-le", "cp1251"):
        try:
            xom = data.decode(kodlash, errors="ignore")
        except (UnicodeDecodeError, LookupError):
            continue
        # Avval TAKRORNI kesamiz — shovqin matnga yopishgan bo'lsa
        # ham matn saqlanib qoladi.
        bolaklar = [_TAKROR_RE.sub(" ", b).strip()
                    for b in _OLE_YAROQLI.findall(xom)]
        bolaklar = [b for b in bolaklar
                    if len(b) >= OLE_MIN_BOLAK and not _ole_shovqinmi(b)]
        matn = "\n".join(bolaklar)
        ball = sum(1 for ch in matn if ch.isalpha())
        if ball > eng_ball:
            eng_matn, eng_ball = matn, ball

    if not eng_matn:
        return "", None, "ole2-xom", "OLE2 ichida o'qiladigan matn topilmadi"
    return clean(eng_matn), None, "ole2-xom", None


def sniff_magic(data: bytes, ext: str) -> str:
    """HAQIQIY formatni baytlardan aniqlaydi; topolmasa `ext` qaytadi.

    ZIP oilasi alohida ajratiladi: `docx` va `xlsx` ham ZIP, farqi
    ichidagi yo'llarda.
    """
    tur = next((n for m, n in _MAGIC if data.startswith(m)), None)
    if tur is None:
        return ext
    if tur != "zip":
        return tur
    try:
        with zipfile.ZipFile(io.BytesIO(data)) as z:
            nomlar = set(z.namelist())
    except (zipfile.BadZipFile, OSError):
        return ext
    if any(n.startswith("word/") for n in nomlar):
        return "docx"
    if any(n.startswith("xl/") for n in nomlar):
        return "xlsx"
    if any(n.startswith("ppt/") for n in nomlar):
        return "pptx"          # parser yo'q, lekin rost nomi bilan rad etamiz
    return "zip"


def _azo_nomi(nom: str) -> str:
    """A'zo nomini YORLIQ sifatida xavfsiz ko'rinishga keltiradi."""
    toza = nom.replace("\\", "/").split("/")[-1]
    return "".join(ch for ch in toza if ch.isprintable())[:120] or "?"


def extract_zip(data: bytes, depth: int = 0) -> Tuple[str, Optional[int], str, Optional[str]]:
    """ZIP ichidagi qo'llab-quvvatlanadigan fayllarni ajratadi.

    Har a'zo matni o'z sarlavhasi bilan qo'shiladi — javobdagi iqtibos
    qaysi fayldan kelganini ko'rsatsin.
    """
    parchalar: List[str] = []
    jami = 0
    ochildi = xato_soni = 0
    try:
        z = zipfile.ZipFile(io.BytesIO(data))
    except (zipfile.BadZipFile, OSError) as e:
        return "", None, "zipfile", f"arxiv ochilmadi: {e}"

    with z:
        azolar = [i for i in z.infolist() if not i.is_dir()]
        if len(azolar) > ZIP_MAX_MEMBERS:
            return ("", None, "zipfile",
                    f"arxivda {len(azolar)} a'zo (chegara {ZIP_MAX_MEMBERS}) "
                    "— ochilmadi")
        for info in azolar:
            # CHEGARALAR — OCHISHDAN OLDIN, e'lon qilingan hajm bo'yicha.
            if info.file_size > ZIP_MAX_MEMBER:
                xato_soni += 1
                continue
            if jami + info.file_size > ZIP_MAX_TOTAL:
                parchalar.append("[...arxiv hajm chegarasida to'xtatildi]")
                break

            ichki_ext = sniff_ext({"name": info.filename})
            try:
                azo = z.read(info)
            except (RuntimeError, zipfile.BadZipFile, OSError):
                xato_soni += 1     # parolli yoki buzilgan a'zo
                continue
            jami += len(azo)

            haqiqiy = sniff_magic(azo, ichki_ext)
            if haqiqiy == "zip":
                if depth >= ZIP_MAX_DEPTH:
                    xato_soni += 1
                    continue
                matn, _, _, xato = extract_zip(azo, depth + 1)
            elif is_supported(haqiqiy):
                matn, _, _, xato = extract(azo, haqiqiy)
            else:
                xato_soni += 1
                continue

            if xato or not matn.strip():
                xato_soni += 1
                continue
            ochildi += 1
            parchalar.append(f"### {_azo_nomi(info.filename)}\n{matn}")

    if not parchalar:
        return ("", None, "zipfile", "arxiv ichida o'qiladigan hujjat yo'q "
                f"({len(azolar)} a'zo, {xato_soni} tasi o'qilmadi)")
    return ("\n\n".join(parchalar), None,
            f"zipfile({ochildi}/{len(azolar)})", None)


def sniff_ext(row: dict) -> str:
    """Kengaytmani aniqlaydi: file_type -> nom -> content_type."""
    ext = (row.get("file_type") or "").strip().lower().lstrip(".")
    if not ext:
        name = row.get("name") or ""
        if "." in name:
            ext = name.rsplit(".", 1)[-1].strip().lower()
    if not ext:
        ct = (row.get("content_type") or "").lower()
        if "pdf" in ct:
            ext = "pdf"
        elif "wordprocessingml" in ct:
            ext = "docx"
        elif "spreadsheetml" in ct:
            ext = "xlsx"
        elif "text/" in ct:
            ext = "txt"
    return ext


def extract(data: bytes, ext: str) -> Tuple[str, Optional[int], str, Optional[str]]:
    if ext in EXT_PDF:
        return extract_pdf(data)
    if ext in EXT_DOCX:
        return extract_docx(data)
    if ext in EXT_XLSX:
        return extract_xlsx(data)
    if ext in EXT_PLAIN:
        return extract_plain(data, ext)
    if ext in EXT_ZIP:
        return extract_zip(data)
    if ext in EXT_OLE2:
        return extract_doc(data)
    return "", None, "", None      # bu yerga kelmasligi kerak


def is_supported(ext: str) -> bool:
    return (ext in EXT_PDF or ext in EXT_DOCX or ext in EXT_XLSX
            or ext in EXT_PLAIN or ext in EXT_ZIP
            or ext in EXT_OLE2)


# ---------------------------------------------------------------------------
# Bitta hujjatni qayta ishlash
# ---------------------------------------------------------------------------
def process(session: requests.Session, row: dict) -> dict:
    """-> tender_document_text ga yoziladigan qator (har doim status bilan)."""
    out = {"tender_id": row["tender_id"], "file_ref": row["file_ref"],
           "text": None, "status": "unreadable", "char_count": None,
           "page_count": None, "error": None, "extractor": None}

    ext = sniff_ext(row)

    # `.doc` — kengaytma bo'yicha rad ETMAYMIZ: haqiqiy turi OLE2
    # bo'lsa `extract_doc` o'qiydi, `docx` bo'lsa (namunada ~8%)
    # oddiy yo'ldan ketadi. Ikkalasi ham baytlarni ko'rgandan keyin
    # aniqlanadi.
    if ext == "doc":
        ext = "ole2"

    # 1. Format — yuklab olishdan OLDIN tekshiramiz (trafik tejaladi)
    if not is_supported(ext):
        out["status"] = "unsupported"
        if not ext:
            out["error"] = "format aniqlanmadi (kengaytma ham, content_type ham yo'q)"
        elif ext in EXT_KNOWN_UNSUPPORTED:
            out["error"] = f"'{ext}' formati qo'llab-quvvatlanmaydi (arxiv/eski binar/rasm)"
        else:
            out["error"] = f"'{ext}' formati uchun parser yo'q"
        return out

    # 2. Hajm — bazadagi metama'lumot bo'yicha
    size = row.get("size_bytes")
    if size and int(size) > MAX_BYTES:
        out["status"] = "too_large"
        out["error"] = f"{int(size) / 1048576:.1f} MB > {MAX_BYTES / 1048576:.0f} MB"
        return out

    # 3. Yuklab olish
    data, err = download(session, row)
    if data is None:
        if err and err.startswith("__TOO_LARGE__"):
            out["status"] = "too_large"
            out["error"] = f"oqimda chegaradan oshdi ({err[13:]} bayt)"
        else:
            out["status"] = "download_failed"
            out["error"] = err
        return out
    if not data:
        out["status"] = "unreadable"
        out["error"] = "bo'sh fayl (0 bayt)"
        return out

    # 4. HAQIQIY format — baytlardan. Kengaytma yolg'on bo'lishi mumkin
    #    (namunada `.doc` fayllarning ~17% i aslida `docx` edi).
    haqiqiy = sniff_magic(data, ext)
    if haqiqiy != ext:
        out["error"] = f"kengaytma '{ext}', haqiqiy format '{haqiqiy}'"
        if not is_supported(haqiqiy):
            out["status"] = "unsupported"
            return out
        ext = haqiqiy

    # 5. Matn ajratish
    text, pages, extractor, perr = extract(data, ext)
    out["extractor"] = extractor or None
    out["page_count"] = pages

    if perr:
        out["status"] = "unreadable"
        out["error"] = perr
        return out
    letters = sum(1 for ch in text if ch.isalpha())
    if len(text) < MIN_CHARS or letters < MIN_LETTERS:
        # Eng ko'p uchraydigan holat: SKAN qilingan PDF yoki CHIZMA — OCR kerak.
        out["status"] = "unreadable"
        out["error"] = (f"matn topilmadi (skan/chizma bo'lishi mumkin, OCR kerak; "
                        f"{len(text)} belgi, {letters} harf)"
                        if ext == "pdf" else
                        f"matn topilmadi (bo'sh hujjat; {len(text)} belgi)")
        return out

    if len(text) > MAX_CHARS:
        text = text[:MAX_CHARS] + "\n\n[...matn chegarada qisqartirildi]"
    out["text"] = text
    out["char_count"] = len(text)
    out["status"] = "ok"
    # Format haqidagi eslatma xato emas — muvaffaqiyatda tozalaymiz.
    if (out.get("error") or "").startswith("kengaytma "):
        out["error"] = None
    return out


# ---------------------------------------------------------------------------
# DB qatlami
# ---------------------------------------------------------------------------
COLS = ["tender_id", "file_ref", "text", "status", "char_count",
        "page_count", "error", "extractor"]


def catalog_tender_ids(conn) -> List[int]:
    """Katalogga mos OCHIQ tenderlar id lari.

    Qoida `api/matching.product_matches()` da — `/catalog/match` endpointi
    bilan AYNAN bir xil (kategoriya roll-up YOKI nom/kalit so'z, alifbodan
    qat'i nazar). Bu yerda takrorlanmaydi (reja_ai_chat.md §15.3.1).

    KO'P KOMPANIYA: `catalog_product` hozir filtrlanmaydi, ya'ni natija
    barcha kompaniyalar kataloglarining BIRLASHMASI — aynan shu kerak,
    aks holda B kompaniyasining tenderlari o'qilmay qolardi (§15.3.2).
    """
    from api import matching, queries        # sof funksiyalar, DB ochmaydi

    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute("SELECT id, name, category_code, keywords FROM catalog_product")
        products = [dict(r) for r in cur.fetchall()]
        if not products:
            # JIMGINA O'TKAZIB YUBORILMAYDI: bo'sh katalog = bo'sh qamrov =
            # hech qanday hujjat o'qilmaydi. Buni ko'rmasdan qolish
            # "ETL ishladi, lekin matn yo'q" degan chalkashlikni beradi.
            print("[!] Katalog BO'SH — `--catalog` qamrovi hech nimani "
                  "tanlamaydi. Katalogni to'ldiring yoki `--no-only-open` / "
                  "`--category` bilan yurgizing.")
            return []
        where, params = queries.build_tender_filters(status="open")
        cur.execute(queries.match_candidates_sql(where, cap=5000), params)
        cands = [dict(r) for r in cur.fetchall()]

    return [c["id"] for c in cands
            if any(matching.product_matches(c, p) for p in products)]


def fetch_targets(conn, args) -> List[dict]:
    """Qayta ishlanadigan hujjatlar. --force bo'lmasa allaqachon
    ishlanganlari CHIQARIB TASHLANADI — takroriy yuklab olish yo'q."""
    where = ["1=1"]
    params: Dict[str, Any] = {}
    if args.tender_id:
        where.append("d.tender_id = %(tid)s")
        params["tid"] = args.tender_id
    if args.platform:
        where.append("d.source_platform = %(plat)s")
        params["plat"] = args.platform
    if args.file_type:
        where.append("lower(d.file_type) = %(ft)s")
        params["ft"] = args.file_type.strip().lower().lstrip(".")
    if not args.force:
        # QAYTA URINISHNI HISOBGA OLADI.
        #
        # Ilgari shart faqat `t.file_ref IS NULL` edi: `download_failed`
        # ham qator qoldirardi, ya'ni bir marta yiqilgan hujjat
        # BOSHQA HECH QACHON olinmasdi. O'tkinchi tarmoq xatosi
        # doimiy nosozlikka aylanardi.
        #
        # Endi: matn qatori yo'q BO'LSA, YOKI yuklab olish yiqilgan
        # va kutish oynasi tugagan bo'lsa olinadi. `butunlay_yiqildi`
        # OLINMAYDI — urinishlar tugagan.
        where.append(
            "(t.file_ref IS NULL "
            " OR (d.holat = 'yuklab_olinmadi' "
            "     AND (d.keyingi_urinish_at IS NULL "
            "          OR d.keyingi_urinish_at <= now())))")
        # Manbada yo'q hujjatni yuklab olishga urinish MA'NOSIZ.
        where.append("d.holat <> 'manbadan_yoqoldi'")

    # --- QAMROV (reja_ai_chat.md §15) --------------------------------------
    # Muddati o'tgan tenderning hujjati qaror uchun kerak emas: bazadagi
    # o'lchov bo'yicha bu filtr yuklab olinadigan hajmni 9.81 GB dan
    # 2.49 GB ga tushiradi va HECH NARSA yo'qotmaydi.
    # `--tender-id` berilgan bo'lsa chetlab o'tiladi: aniq so'ralgan tender
    # muddati o'tgan bo'lsa ham o'qilishi kerak.
    if getattr(args, "only_open", True) and not args.tender_id:
        where.append("EXISTS (SELECT 1 FROM tender t2 WHERE t2.id = d.tender_id "
                     "AND t2.status = 'open' "
                     "AND (t2.close_at IS NULL OR t2.close_at > now()))")

    if getattr(args, "category", None):
        # Parent tanlansa ichkilari ham: 'qurilish' -> 'qurilish/yol'
        where.append("EXISTS (SELECT 1 FROM tender_category tc "
                     "WHERE tc.tender_id = d.tender_id "
                     "AND (tc.code = %(cat)s OR tc.code LIKE %(cat)s || '/%%'))")
        params["cat"] = args.category.strip()

    if getattr(args, "catalog", False):
        ids = catalog_tender_ids(conn)
        if not ids:
            return []                      # katalog bo'sh -> qamrov ham bo'sh
        where.append("d.tender_id = ANY(%(cat_ids)s)")
        params["cat_ids"] = ids

    # ------------------------------------------------------------------
    # TARTIB: PLATFORMALAR ARALASHTIRILADI, KETMA-KET EMAS
    # ------------------------------------------------------------------
    # O'LCHANGAN NUQSON (2026-09-03). Ilgari bu yerda shunchaki
    # `ORDER BY d.tender_id DESC` turardi. Ikki platformaning ID
    # FAZOLARI esa butunlay boshqa:
    #
    #     xt-xarid      108 .. 8 538 264
    #     uzex   20 000 475 229 .. 20 000 510 026
    #
    # Ya'ni HAR uzex hujjati HAR xt-xarid hujjatidan oldin turardi.
    # O'lchandi — navbatdagi o'rinlar:
    #
    #     uzex       1 .. 342   (342 ta)
    #     xt-xarid 343 .. 906   (564 ta)
    #
    # Qadamda vaqt byudjeti bor (to'liq qamrov ~25 daqiqa), shuning
    # uchun BIRORTA xt-xarid hujjati 342 ta uzex hujjati tugamaguncha
    # yetib borilmasdi. Natija: xt-xarid matn qamrovi 31/595, uzex esa
    # 2 925. Bu "sekin ishlayapti" emas — STRUKTURAVIY OCHLIK: filtr
    # ham, byudjet ham to'g'ri, faqat NAVBAT TARTIBI bir platformani
    # abadiy oxirga surardi.
    #
    # Yechim: har platforma ichida tartib SAQLANADI (yangi tender
    # oldin), lekin platformalar NAVBATMA-NAVBAT olinadi. Bu yon
    # foyda ham beradi — ikki HOSTGA so'rov navbatma-navbat ketadi,
    # bitta manbaga ketma-ket urish o'rniga.
    #
    # `--platform` berilgan bo'lsa aralashtirish ma'nosiz, lekin
    # zarari ham yo'q: bitta guruh qoladi.
    sql = f"""
        SELECT tender_id, file_ref, file_id, file_path, name,
               size_bytes, content_type, file_type, source_platform,
               holat, urinish
        FROM (
            SELECT d.tender_id, d.file_ref, d.file_id, d.file_path, d.name,
                   d.size_bytes, d.content_type, d.file_type,
                   d.source_platform, d.holat, d.urinish,
                   row_number() OVER (PARTITION BY d.source_platform
                                      ORDER BY d.tender_id DESC, d.file_ref)
                       AS _navbat
            FROM tender_document d
            LEFT JOIN tender_document_text t
                   ON t.tender_id = d.tender_id AND t.file_ref = d.file_ref
            WHERE {' AND '.join(where)}
        ) s
        ORDER BY _navbat, source_platform, tender_id DESC, file_ref
    """
    if args.limit:
        sql += f" LIMIT {int(args.limit)}"

    with conn.cursor(cursor_factory=RealDictCursor) as cur:
        cur.execute(sql, params)
        return [dict(r) for r in cur.fetchall()]


#: `tender_document_text.status` -> `tender_document.holat`.
#: Ikki lug'at ATAYLAB alohida: birinchisi AJRATISH natijasi,
#: ikkinchisi QAYTA ISHLASH BOSQICHI (unda "navbatda",
#: "rejalashtirilmagan" kabi holatlar ham bor).
HOLAT_XARITA = {
    "ok": "ok",
    "unreadable": "unreadable",
    "unsupported": "unsupported",
    "too_large": "too_large",
    "download_failed": "yuklab_olinmadi",
}

#: Yuklab olish necha marta qayta urinilsin. Chegaradan oshgach holat
#: `butunlay_yiqildi` bo'ladi va hujjat navbatdan CHIQADI.
#:
#: ILGARI QAYTA URINISH UMUMAN YO'Q EDI: `fetch_targets()` "matn
#: qatori bor" degan shartga qarardi, `download_failed` ham qator
#: qoldirardi — ya'ni bir marta yiqilgan hujjat BOSHQA HECH QACHON
#: olinmasdi (`--force` siz). O'tkinchi tarmoq xatosi doimiy
#: nosozlikka aylanardi.
MAX_URINISH = 4


def belgila_boshlandi(conn, row: dict) -> None:
    """Hujjat ustida ISH BOSHLANDI — vaqt belgilari yoziladi.

    Bu qator `process()` DAN OLDIN chaqiriladi. Jarayon o'rtada
    o'ldirilsa holat `yuklanmoqda` bo'lib qoladi va bu HALOL:
    "boshlandi, tugamadi" — "hech qachon urinilmagan" dan boshqa gap.
    """
    with conn.cursor() as cur:
        cur.execute(
            "UPDATE tender_document "
            "SET holat = 'yuklanmoqda', download_started_at = now(), "
            "    extraction_started_at = NULL, extraction_finished_at = NULL "
            "WHERE tender_id = %s AND file_ref = %s",
            (row["tender_id"], row["file_ref"]))
    conn.commit()


def save(conn, rec: dict) -> None:
    """Bitta yozuv (idempotent). Har hujjatdan keyin commit — uzun yurish
    yarmida uzilsa ham ishlangan qism saqlanib qoladi.

    METADATA QATORINING HOLATI HAM YANGILANADI. Ilgari holat FAQAT
    `tender_document_text` da yashardi va "hali urinilmagan" holatining
    bazada KO'RINISHI YO'Q edi — aynan shu 7 603 hujjatni ko'rinmas
    qilgan (`schema_patch_doc_qamrov.sql`).
    """
    holat = HOLAT_XARITA.get(rec["status"], "unreadable")
    xato = rec.get("error")
    with conn.cursor() as cur:
        cur.execute(
            f"""INSERT INTO tender_document_text ({','.join(COLS)})
                VALUES ({','.join('%s' for _ in COLS)})
                ON CONFLICT (tender_id, file_ref) DO UPDATE SET
                {','.join(f'{c}=EXCLUDED.{c}' for c in COLS if c not in ('tender_id', 'file_ref'))},
                extracted_at = now()""",
            [rec[c] for c in COLS])
        # Yuklab olish yiqilgan bo'lsa urinish sanaladi va chegaradan
        # oshgach hujjat navbatdan CHIQADI (`butunlay_yiqildi`).
        # Muvaffaqiyatda sanoq NOLGA qaytadi.
        cur.execute(
            "UPDATE tender_document SET "
            "  urinish = CASE WHEN %(h)s = 'yuklab_olinmadi' "
            "                 THEN urinish + 1 ELSE 0 END, "
            "  holat = CASE WHEN %(h)s = 'yuklab_olinmadi' "
            "                    AND urinish + 1 >= %(max)s "
            "               THEN 'butunlay_yiqildi' ELSE %(h)s END, "
            "  keyingi_urinish_at = CASE WHEN %(h)s = 'yuklab_olinmadi' "
            "                                 AND urinish + 1 < %(max)s "
            "                            THEN now() + (power(2, urinish + 1)"
            "                                          * interval '1 minute') END, "
            "  downloaded_at = CASE WHEN %(h)s IN ('ok','unreadable') "
            "                       THEN now() ELSE downloaded_at END, "
            "  extraction_started_at = COALESCE(extraction_started_at, now()), "
            "  extraction_finished_at = now(), "
            "  last_error = %(e)s, "
            "  last_error_at = CASE WHEN %(e)s IS NULL THEN NULL ELSE now() END "
            "WHERE tender_id = %(t)s AND file_ref = %(f)s",
            {"h": holat, "e": xato, "max": MAX_URINISH,
             "t": rec["tender_id"], "f": rec["file_ref"]})
    conn.commit()


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------
def main() -> None:
    if load_dotenv:
        load_dotenv()

    ap = argparse.ArgumentParser(description="Tender hujjatlaridan matn ajratish ETL")
    ap.add_argument("--limit", type=int, help="Nechta hujjat (sinov uchun)")
    ap.add_argument("--tender-id", type=int, help="Faqat shu tender hujjatlari")
    ap.add_argument("--platform", help="Faqat shu manba ('xt-xarid' | 'uzex')")
    ap.add_argument("--file-type", help="Faqat shu tur ('pdf' | 'docx' | 'xlsx' ...)")
    # --- QAMROV (reja_ai_chat.md §15.4) ---
    ap.add_argument("--only-open", action=argparse.BooleanOptionalAction, default=True,
                    help="Faqat ochiq va muddati tugamagan tenderlar "
                         "(STANDART: yoqilgan; o'chirish: --no-only-open)")
    ap.add_argument("--catalog", action="store_true",
                    help="Faqat katalogga mos tenderlar "
                         "(qoida: api/matching.product_matches)")
    ap.add_argument("--category", help="Faqat shu kategoriya ('elektr', "
                                       "'qurilish' — ichkilari ham kiradi)")
    ap.add_argument("--count-only", action="store_true",
                    help="Faqat SANAYDI va chiqadi — hech narsa yuklab olinmaydi")
    ap.add_argument("--force", action="store_true",
                    help="Allaqachon ishlanganlarni ham qayta yuklab oladi")
    ap.add_argument("--dry-run", action="store_true",
                    help="Yuklab oladi va matn ajratadi, lekin DBga yozmaydi")
    ap.add_argument("--max-seconds", type=float, default=STANDART_BYUDJET,
                    help="Vaqt byudjeti. Tugaganda TOZA to'xtaydi va qolgani keyingi yurishga qoladi (0 = cheksiz)")
    ap.add_argument("--quiet", action="store_true", help="Har fayl uchun satr chiqarmaydi")
    ap.add_argument("--dsn", default=os.environ.get("XT_DB_DSN"))
    args = ap.parse_args()

    global _TOXTATGICH
    _TOXTATGICH = ish.Toxtatgich(getattr(args, 'max_seconds', 0) or None)
    _TOXTATGICH.signallarni_ulash()

    if not args.dsn:
        sys.exit("XATO: DSN yo'q. --dsn yoki XT_DB_DSN o'rnating.")
    if psycopg2 is None:
        sys.exit("XATO: pip install psycopg2-binary")

    conn = psycopg2.connect(args.dsn)
    rows = fetch_targets(conn, args)

    if args.count_only:
        # QAMROVNI KO'RSATADI, TARMOQQA CHIQMAYDI. Yuklab olishdan oldin
        # "nima tanlandi va qancha bo'ladi" degan savolga javob (§15.4).
        scope = []
        if args.only_open and not args.tender_id:
            scope.append("ochiq + muddati tugamagan")
        if args.catalog:
            scope.append("katalogga mos")
        if args.category:
            scope.append(f"kategoriya={args.category}")
        if args.platform:
            scope.append(f"platforma={args.platform}")
        if args.file_type:
            scope.append(f"tur={args.file_type}")
        print(f"Qamrov: {' · '.join(scope) or 'cheklovsiz'}")

        by_type: Dict[str, List[int]] = {}
        sup_n = sup_b = uns_n = 0
        for r in rows:
            ext = sniff_ext(r)
            b = r.get("size_bytes") or 0
            acc = by_type.setdefault(ext or "?", [0, 0])
            acc[0] += 1
            acc[1] += b
            if is_supported(ext):
                sup_n += 1
                sup_b += b
            else:
                uns_n += 1

        for ext, (n, b) in sorted(by_type.items(), key=lambda x: -x[1][1])[:12]:
            mark = "" if is_supported(ext) else "  (qo'llab-quvvatlanmaydi)"
            print(f"  {ext:<8} {n:>5} ta  {b / 1048576:>8.1f} MB{mark}")

        print(f"\n  Matn ajratiladi: {sup_n} ta, {sup_b / 1073741824:.2f} GB")
        print(f"  O'tkazib yuboriladi: {uns_n} ta")
        print(f"  Taxminiy vaqt: ~{sup_n * (REQUEST_DELAY + 1.2) / 60:.0f} daqiqa")
        print("\n  (--count-only — hech narsa yuklab olinmadi)")
        conn.close()
        return

    if not rows:
        print("Qayta ishlanadigan hujjat yo'q "
              "(hammasi ishlangan bo'lishi mumkin — --force bilan qayta yurgizing).")
        conn.close()
        return

    print(f"[1/2] {len(rows)} ta hujjat qayta ishlanadi"
          f"{' (--force)' if args.force else ''}...")
    print(f"      Taxminiy vaqt: ~{len(rows) * (REQUEST_DELAY + 1.2) / 60:.1f} daqiqa\n")

    session = requests.Session()
    chiqish_qismam = False
    counts: Dict[str, int] = {}
    total_chars = 0
    downloaded = 0

    for i, row in enumerate(rows, 1):
        # VAQT BYUDJETI / TO'XTASH SO'ROVI — TOZA CHIQISH.
        #
        # NEGA QO'SHILDI (2026-09-03). Bu quvurdagi ENG SEKIN qadam
        # (to'liq qamrov ~25 daqiqa), lekin `etl_uzex.py` va
        # `etl_details.py` dan farqli o'laroq unda byudjet YO'Q edi.
        # Oqibati o'lchangan: `LastTaskResult=0xC000013A`
        # (STATUS_CONTROL_C_EXIT) — host CTRL+C yuborganda qadam
        # FAYL O'RTASIDA o'lardi.
        #
        # `ish.Toxtatgich` SIGINT/SIGBREAK/SIGTERM ni ham ushlaydi,
        # ya'ni o'sha CTRL+C endi O'LIM emas, TOZA TO'XTASH bo'ladi:
        # joriy fayl tugatiladi, qolgani keyingi yurishga qoladi.
        # Navbatning o'zi checkpoint vazifasini bajaradi — ishlangan
        # hujjat `tender_document_text` da qator qoldiradi va
        # `fetch_targets()` uni boshqa tanlamaydi.
        if _TOXTATGICH is not None and _TOXTATGICH.toxtaymi():
            sabab = _TOXTATGICH.sabab or "toxtatildi"
            chiqish_qismam = True
            print(f"\n[!] TO'XTASH ({sabab}): {i - 1}/{len(rows)} bajarildi. "
                  "Navbat saqlanadi — keyingi yurish qolganidan davom etadi.")
            break

        # ISH BOSHLANGANINI DARHOL BELGILAYMIZ. Jarayon o'rtada
        # o'ldirilsa holat `yuklanmoqda` bo'lib qoladi — bu HALOL
        # ("boshlandi, tugamadi") va keyingi yurish uni qayta oladi.
        if not args.dry_run:
            try:
                belgila_boshlandi(conn, row)
            except Exception as e:                          # noqa: BLE001
                conn.rollback()
                print(f"    ! holat belgilanmadi: {str(e)[:90]}", file=sys.stderr)
        rec = process(session, row)
        counts[rec["status"]] = counts.get(rec["status"], 0) + 1
        total_chars += rec["char_count"] or 0

        if not args.quiet:
            name = (row.get("name") or row["file_ref"])[:48]
            extra = (f"{rec['char_count']} belgi" if rec["status"] == "ok"
                     else (rec["error"] or "")[:60])
            print(f"  [{i}/{len(rows)}] #{row['tender_id']} {name} "
                  f"-> {rec['status']} ({extra})")

        if not args.dry_run:
            try:
                save(conn, rec)
            except Exception as e:  # noqa: BLE001
                conn.rollback()
                print(f"    ! DB xato: {e}", file=sys.stderr)

        # Kechikish FAQAT tarmoqqa chiqqan bo'lsak (unsupported/too_large tez o'tadi)
        if rec["status"] not in ("unsupported", "too_large"):
            downloaded += 1
            time.sleep(REQUEST_DELAY)

    conn.close()

    print(f"\n[2/2] Tayyor. Yuklab olingan fayl: {downloaded}, "
          f"jami matn: {total_chars:,} belgi")
    for st in sorted(counts, key=lambda k: -counts[k]):
        print(f"      {st:<16} {counts[st]}")
    manual = sum(v for k, v in counts.items() if k != "ok")
    print(f"      -> qo'lda tekshirish talab etiladi: {manual}")
    if args.dry_run:
        print("      (--dry-run — DBga yozilmadi)")

    # QISMAN tugash XATO EMAS va MUVAFFAQIYAT ham emas. `run_etl.py`
    # `7` ni `partial` deb o'qiydi va navbat keyingi yurishda davom
    # etadi — "hammasi bajarildi" degan YOLG'ON chiqmaydi.
    if chiqish_qismam:
        sys.exit(CHIQISH_QISMAN)


if __name__ == "__main__":
    main()
